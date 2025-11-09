from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_resume(filename, name, email, education, experience, skills, projects):
    """DOCX 이력서 생성"""
    doc = Document()

    # 이름 (제목)
    title = doc.add_heading(name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 연락처
    contact = doc.add_paragraph(f"Email: {email}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 학력
    doc.add_heading("학력", level=1)
    doc.add_paragraph(education)

    # 경력
    doc.add_heading("경력", level=1)
    for exp in experience:
        doc.add_paragraph(exp, style="List Bullet")

    # 기술 스택
    doc.add_heading("기술 스택", level=1)
    doc.add_paragraph(skills)

    # 프로젝트
    doc.add_heading("주요 프로젝트", level=1)
    for project in projects:
        doc.add_paragraph(project, style="List Bullet")

    # 저장
    doc.save(filename)
    print(f"Created: {filename}")


# 이력서 데이터 (30명)
resumes = [
    # Backend Engineers (FastAPI/Django) - 10명
    {
        "filename": "data/resumes/김철수_이력서.docx",
        "name": "김철수",
        "email": "chulsoo.kim@email.com",
        "education": "서울대학교 컴퓨터공학과 학사 (2015-2019)",
        "experience": [
            "네이버 AI Lab - ML Engineer (2021-2024, 3년)",
            "카카오 - Backend Developer (2019-2021, 2년)",
        ],
        "skills": "Python, FastAPI, Django, PostgreSQL, MongoDB, Docker, Kubernetes, AWS, Sentence Transformers, ChromaDB, LangChain, OpenAI API",
        "projects": [
            "AI 챗봇 시스템 구축 - FastAPI + LangChain + ChromaDB를 활용한 RAG 기반 챗봇 개발",
            "추천 시스템 API - Sentence Transformers를 활용한 의미 기반 상품 추천 시스템",
            "대용량 로그 처리 파이프라인 - Kafka + Spark를 활용한 실시간 로그 분석 시스템",
        ],
    },
    {
        "filename": "data/resumes/이영희_이력서.docx",
        "name": "이영희",
        "email": "younghee.lee@email.com",
        "education": "KAIST 전산학부 석사 (2018-2020), 고려대 컴퓨터학과 학사 (2014-2018)",
        "experience": [
            "쿠팡 - Senior Backend Engineer (2020-2024, 4년)",
        ],
        "skills": "Python, FastAPI, Flask, PostgreSQL, Redis, Docker, AWS, GCP, Vector DB (Pinecone), Hugging Face Transformers",
        "projects": [
            "상품 검색 시스템 개선 - Vector DB를 활용한 의미 기반 검색 시스템 구축",
            "ML 모델 서빙 플랫폼 - FastAPI 기반 모델 서빙 인프라 구축, 일 1억건 이상 처리",
            "실시간 추천 엔진 - Redis + ML 모델을 활용한 실시간 개인화 추천 시스템",
        ],
    },
    {
        "filename": "data/resumes/박민수_이력서.docx",
        "name": "박민수",
        "email": "minsu.park@email.com",
        "education": "연세대학교 소프트웨어학과 학사 (2017-2021)",
        "experience": [
            "스타트업 A - Backend Developer (2021-2024, 3년)",
        ],
        "skills": "Python, FastAPI, Django, MySQL, MongoDB, Git, Docker",
        "projects": [
            "RESTful API 개발 - FastAPI를 활용한 B2B SaaS 플랫폼 백엔드 개발",
            "데이터베이스 최적화 - 쿼리 최적화를 통한 API 응답속도 50% 개선",
            "사용자 인증 시스템 - JWT 기반 인증/인가 시스템 구축",
        ],
    },
    {
        "filename": "data/resumes/정수진_이력서.docx",
        "name": "정수진",
        "email": "soojin.jung@email.com",
        "education": "부산대학교 컴퓨터공학과 학사 (2019-2023)",
        "experience": [
            "IT 스타트업 - Junior Backend Developer (2023-2024, 1년 6개월)",
        ],
        "skills": "Python, Flask, Django, MySQL, Git",
        "projects": [
            "사내 관리 시스템 개발 - Django를 활용한 사내 업무 관리 시스템",
            "간단한 REST API 구현 - Flask 기반 CRUD API 개발",
        ],
    },
    {
        "filename": "data/resumes/강지훈_이력서.docx",
        "name": "강지훈",
        "email": "jihoon.kang@email.com",
        "education": "성균관대학교 컴퓨터공학과 학사 (2012-2016)",
        "experience": [
            "라인 - Lead Backend Engineer (2020-2024, 4년)",
            "우아한형제들 - Senior Backend Developer (2016-2020, 4년)",
        ],
        "skills": "Python, FastAPI, Django, PostgreSQL, Redis, RabbitMQ, Celery, Docker, Kubernetes, AWS, Terraform",
        "projects": [
            "마이크로서비스 아키텍처 구축 - FastAPI 기반 MSA 전환 프로젝트 리딩",
            "비동기 작업 처리 시스템 - Celery + RabbitMQ를 활용한 대용량 백그라운드 작업 처리",
            "API Gateway 개발 - 인증/인가, Rate Limiting, 로깅을 포함한 API Gateway 구축",
        ],
    },
    {
        "filename": "data/resumes/윤서연_이력서.docx",
        "name": "윤서연",
        "email": "seoyeon.yoon@email.com",
        "education": "이화여대 컴퓨터공학과 학사 (2018-2022)",
        "experience": [
            "토스 - Backend Developer (2022-2024, 2년)",
        ],
        "skills": "Python, FastAPI, PostgreSQL, Redis, Docker, AWS, GraphQL",
        "projects": [
            "결제 API 개발 - FastAPI를 활용한 고가용성 결제 시스템 구축",
            "GraphQL API 서버 구축 - Strawberry를 활용한 GraphQL 서버 개발",
            "캐싱 전략 최적화 - Redis를 활용한 API 성능 3배 향상",
        ],
    },
    {
        "filename": "data/resumes/송민재_이력서.docx",
        "name": "송민재",
        "email": "minjae.song@email.com",
        "education": "한국과학기술원(KAIST) 전산학부 학사 (2016-2020)",
        "experience": [
            "당근마켓 - Backend Engineer (2020-2024, 4년)",
        ],
        "skills": "Python, Django, FastAPI, PostgreSQL, MySQL, MongoDB, Redis, Kafka, Docker, Kubernetes, AWS",
        "projects": [
            "실시간 채팅 시스템 - WebSocket + Redis Pub/Sub을 활용한 대규모 채팅 서비스",
            "이벤트 드리븐 아키텍처 구축 - Kafka를 활용한 MSA 간 이벤트 처리",
            "검색 시스템 개선 - Elasticsearch를 활용한 지역 기반 검색 최적화",
        ],
    },
    {
        "filename": "data/resumes/한소희_이력서.docx",
        "name": "한소희",
        "email": "sohee.han@email.com",
        "education": "서강대학교 컴퓨터공학과 학사 (2020-2024)",
        "experience": [
            "핀테크 스타트업 - Junior Backend Developer (2024-현재, 6개월)",
        ],
        "skills": "Python, FastAPI, PostgreSQL, Docker, Git, REST API",
        "projects": [
            "사용자 관리 API - FastAPI를 활용한 회원가입/로그인 API 개발",
            "데이터 마이그레이션 - 레거시 시스템에서 신규 DB로 데이터 이전",
        ],
    },
    {
        "filename": "data/resumes/오태양_이력서.docx",
        "name": "오태양",
        "email": "taeyang.oh@email.com",
        "education": "중앙대학교 소프트웨어학과 학사 (2013-2017)",
        "experience": [
            "SK텔레콤 - Senior Backend Engineer (2017-2024, 7년)",
        ],
        "skills": "Python, Django, FastAPI, PostgreSQL, Redis, RabbitMQ, Docker, Kubernetes, AWS, Azure",
        "projects": [
            "통신사 빌링 시스템 - Django 기반 대용량 과금 시스템 개발 및 운영",
            "고객 포털 API 서버 - FastAPI를 활용한 RESTful API 서버 개발",
            "레거시 시스템 모더나이제이션 - 모놀리식 → MSA 전환 프로젝트",
        ],
    },
    {
        "filename": "data/resumes/임하늘_이력서.docx",
        "name": "임하늘",
        "email": "haneul.lim@email.com",
        "education": "포항공대(POSTECH) 컴퓨터공학과 석사 (2017-2019), 학사 (2013-2017)",
        "experience": [
            "삼성SDS - Principal Engineer (2019-2024, 5년)",
        ],
        "skills": "Python, FastAPI, Django, Go, PostgreSQL, MongoDB, Redis, Kafka, Docker, Kubernetes, AWS, GCP, Terraform, Istio",
        "projects": [
            "클라우드 네이티브 플랫폼 구축 - Kubernetes 기반 사내 PaaS 플랫폼 설계 및 개발",
            "MSA 프레임워크 개발 - FastAPI 기반 사내 표준 마이크로서비스 프레임워크",
            "API Gateway 및 서비스 메시 - Istio를 활용한 트래픽 관리 시스템",
        ],
    },

    # ML/AI Engineers - 5명
    {
        "filename": "data/resumes/최지원_이력서.docx",
        "name": "최지원",
        "email": "jiwon.choi@email.com",
        "education": "서울대학교 컴퓨터공학과 박사 (2018-2023), 석사 (2016-2018), 학사 (2012-2016)",
        "experience": [
            "네이버 클로바 - Senior AI Research Engineer (2023-2024, 1년)",
        ],
        "skills": "Python, PyTorch, TensorFlow, Hugging Face Transformers, FastAPI, LangChain, Vector DB, AWS SageMaker, MLflow",
        "projects": [
            "LLM 파인튜닝 - 한국어 도메인 특화 LLM 파인튜닝 및 배포",
            "RAG 시스템 구축 - LangChain + ChromaDB를 활용한 엔터프라이즈 RAG 시스템",
            "모델 서빙 인프라 - FastAPI 기반 ML 모델 서빙 파이프라인 구축",
        ],
    },
    {
        "filename": "data/resumes/배준호_이력서.docx",
        "name": "배준호",
        "email": "junho.bae@email.com",
        "education": "KAIST AI대학원 석사 (2020-2022), 성균관대 소프트웨어학과 학사 (2016-2020)",
        "experience": [
            "카카오브레인 - ML Engineer (2022-2024, 2년)",
        ],
        "skills": "Python, PyTorch, FastAPI, Docker, Kubernetes, Sentence Transformers, FAISS, Pinecone, MLOps",
        "projects": [
            "이미지 검색 시스템 - CLIP 모델을 활용한 멀티모달 검색 엔진",
            "추천 시스템 - 협업 필터링 및 딥러닝 기반 개인화 추천",
            "ML 파이프라인 자동화 - Kubeflow를 활용한 MLOps 파이프라인",
        ],
    },
    {
        "filename": "data/resumes/신예린_이력서.docx",
        "name": "신예린",
        "email": "yerin.shin@email.com",
        "education": "연세대학교 인공지능학과 석사 (2021-2023), 학사 (2017-2021)",
        "experience": [
            "업스테이지 - AI Engineer (2023-2024, 1년)",
        ],
        "skills": "Python, PyTorch, FastAPI, LangChain, OpenAI API, Anthropic API, Pinecone, Streamlit",
        "projects": [
            "문서 QA 시스템 - LangChain을 활용한 법률 문서 질의응답 시스템",
            "프롬프트 엔지니어링 플랫폼 - GPT-4를 활용한 자동 프롬프트 최적화",
            "AI 챗봇 서비스 - RAG 기반 고객 상담 챗봇 개발",
        ],
    },
    {
        "filename": "data/resumes/권도현_이력서.docx",
        "name": "권도현",
        "email": "dohyun.kwon@email.com",
        "education": "고려대학교 데이터과학과 학사 (2019-2023)",
        "experience": [
            "AI 스타트업 - Junior ML Engineer (2023-2024, 1년)",
        ],
        "skills": "Python, scikit-learn, PyTorch, FastAPI, Pandas, NumPy, Jupyter",
        "projects": [
            "수요 예측 모델 - 시계열 분석을 활용한 재고 최적화",
            "감성 분석 API - KoBERT를 활용한 리뷰 감성 분석 서비스",
        ],
    },
    {
        "filename": "data/resumes/장서현_이력서.docx",
        "name": "장서현",
        "email": "seohyun.jang@email.com",
        "education": "MIT EECS 석사 (2019-2021), 서울대학교 컴퓨터공학과 학사 (2015-2019)",
        "experience": [
            "OpenAI - Research Engineer (2021-2023, 2년)",
            "구글 - Software Engineer (2023-2024, 1년)",
        ],
        "skills": "Python, PyTorch, JAX, C++, FastAPI, Distributed Training, RLHF, Model Compression, TensorRT",
        "projects": [
            "LLM 학습 최적화 - RLHF를 활용한 대규모 언어모델 파인튜닝",
            "모델 경량화 - Quantization 및 Distillation을 통한 모델 압축",
            "분산 학습 시스템 - Multi-GPU/Multi-Node 학습 파이프라인 구축",
        ],
    },

    # DevOps/Infrastructure - 3명
    {
        "filename": "data/resumes/조민석_이력서.docx",
        "name": "조민석",
        "email": "minseok.jo@email.com",
        "education": "한양대학교 컴퓨터공학과 학사 (2015-2019)",
        "experience": [
            "쿠팡 - Senior DevOps Engineer (2019-2024, 5년)",
        ],
        "skills": "Kubernetes, Docker, Terraform, Ansible, Jenkins, GitLab CI, AWS, GCP, Prometheus, Grafana, ELK Stack",
        "projects": [
            "CI/CD 파이프라인 구축 - GitLab CI + ArgoCD를 활용한 자동화 배포",
            "인프라 코드화(IaC) - Terraform을 활용한 멀티 클라우드 인프라 관리",
            "모니터링 시스템 - Prometheus + Grafana를 활용한 실시간 모니터링",
        ],
    },
    {
        "filename": "data/resumes/홍지수_이력서.docx",
        "name": "홍지수",
        "email": "jisu.hong@email.com",
        "education": "부산대학교 정보컴퓨터공학과 학사 (2017-2021)",
        "experience": [
            "네이버 - DevOps Engineer (2021-2024, 3년)",
        ],
        "skills": "Docker, Kubernetes, Jenkins, ArgoCD, AWS, Terraform, Python, Bash, Prometheus, Grafana",
        "projects": [
            "Kubernetes 클러스터 운영 - 100+ 마이크로서비스 운영 및 관리",
            "자동화 스크립트 개발 - Python/Bash를 활용한 운영 자동화",
            "장애 대응 시스템 - PagerDuty + Slack 연동 알림 시스템",
        ],
    },
    {
        "filename": "data/resumes/서준우_이력서.docx",
        "name": "서준우",
        "email": "junwoo.seo@email.com",
        "education": "경북대학교 컴퓨터학과 학사 (2018-2022)",
        "experience": [
            "카카오 - Junior DevOps Engineer (2022-2024, 2년)",
        ],
        "skills": "Docker, Kubernetes, Jenkins, AWS, Linux, Git, Python",
        "projects": [
            "컨테이너 환경 구축 - Docker 기반 개발 환경 표준화",
            "배포 자동화 - Jenkins를 활용한 CI/CD 파이프라인 구축",
        ],
    },

    # Data Engineers - 3명
    {
        "filename": "data/resumes/안유진_이력서.docx",
        "name": "안유진",
        "email": "yujin.ahn@email.com",
        "education": "서울대학교 통계학과 석사 (2018-2020), 학사 (2014-2018)",
        "experience": [
            "11번가 - Senior Data Engineer (2020-2024, 4년)",
        ],
        "skills": "Python, Apache Spark, Airflow, Kafka, Hadoop, AWS EMR, Redshift, PostgreSQL, dbt",
        "projects": [
            "실시간 데이터 파이프라인 - Kafka + Spark Streaming을 활용한 실시간 분석",
            "데이터 웨어하우스 구축 - Redshift 기반 DWH 설계 및 ETL 파이프라인",
            "데이터 품질 관리 - Great Expectations를 활용한 데이터 검증 시스템",
        ],
    },
    {
        "filename": "data/resumes/이동욱_이력서.docx",
        "name": "이동욱",
        "email": "dongwook.lee@email.com",
        "education": "KAIST 전산학부 학사 (2016-2020)",
        "experience": [
            "쿠팡 - Data Engineer (2020-2024, 4년)",
        ],
        "skills": "Python, Scala, Spark, Airflow, Kafka, Flink, AWS, Snowflake, dbt, Looker",
        "projects": [
            "대용량 데이터 처리 - Spark를 활용한 일 10TB+ 데이터 처리 파이프라인",
            "스트리밍 데이터 처리 - Flink를 활용한 실시간 이벤트 처리",
            "데이터 마트 구축 - dbt를 활용한 분석용 데이터 마트 개발",
        ],
    },
    {
        "filename": "data/resumes/박채원_이력서.docx",
        "name": "박채원",
        "email": "chaewon.park@email.com",
        "education": "연세대학교 응용통계학과 학사 (2019-2023)",
        "experience": [
            "데이터 스타트업 - Junior Data Engineer (2023-2024, 1년)",
        ],
        "skills": "Python, SQL, Airflow, PostgreSQL, BigQuery, Pandas, dbt",
        "projects": [
            "ETL 파이프라인 개발 - Airflow를 활용한 일배치 데이터 처리",
            "데이터 분석 자동화 - Python을 활용한 리포트 자동 생성",
        ],
    },

    # Frontend/Fullstack - 5명
    {
        "filename": "data/resumes/김태희_이력서.docx",
        "name": "김태희",
        "email": "taehee.kim@email.com",
        "education": "이화여대 소프트웨어학과 학사 (2016-2020)",
        "experience": [
            "토스 - Senior Frontend Engineer (2020-2024, 4년)",
        ],
        "skills": "TypeScript, React, Next.js, Vue.js, Tailwind CSS, Zustand, React Query, Webpack, Vite",
        "projects": [
            "디자인 시스템 구축 - React 기반 사내 컴포넌트 라이브러리 개발",
            "웹 성능 최적화 - Core Web Vitals 개선을 통한 LCP 50% 향상",
            "SSR/SSG 적용 - Next.js를 활용한 SEO 최적화",
        ],
    },
    {
        "filename": "data/resumes/정현우_이력서.docx",
        "name": "정현우",
        "email": "hyunwoo.jung@email.com",
        "education": "고려대학교 컴퓨터학과 학사 (2017-2021)",
        "experience": [
            "당근마켓 - Fullstack Engineer (2021-2024, 3년)",
        ],
        "skills": "TypeScript, React, Next.js, Node.js, NestJS, PostgreSQL, Redis, Docker, AWS",
        "projects": [
            "실시간 알림 시스템 - WebSocket을 활용한 실시간 푸시 알림",
            "풀스택 채팅 기능 - React + NestJS 기반 실시간 채팅 개발",
            "이미지 업로드 최적화 - S3 + CloudFront를 활용한 이미지 처리",
        ],
    },
    {
        "filename": "data/resumes/이수빈_이력서.docx",
        "name": "이수빈",
        "email": "subin.lee@email.com",
        "education": "성균관대학교 소프트웨어학과 학사 (2018-2022)",
        "experience": [
            "라인 - Frontend Developer (2022-2024, 2년)",
        ],
        "skills": "JavaScript, TypeScript, React, Redux, Styled-components, Jest, Storybook",
        "projects": [
            "관리자 대시보드 개발 - React를 활용한 데이터 시각화 대시보드",
            "컴포넌트 테스트 자동화 - Jest + React Testing Library 도입",
            "UI/UX 개선 - 사용자 피드백 기반 인터페이스 개선",
        ],
    },
    {
        "filename": "data/resumes/최윤서_이력서.docx",
        "name": "최윤서",
        "email": "yunseo.choi@email.com",
        "education": "중앙대학교 소프트웨어학과 학사 (2019-2023)",
        "experience": [
            "프론트엔드 스타트업 - Junior Frontend Developer (2023-2024, 1년)",
        ],
        "skills": "JavaScript, React, HTML, CSS, Git",
        "projects": [
            "반응형 웹 개발 - React를 활용한 모바일 최적화 웹사이트",
            "UI 컴포넌트 개발 - 재사용 가능한 React 컴포넌트 라이브러리",
        ],
    },
    {
        "filename": "data/resumes/강민지_이력서.docx",
        "name": "강민지",
        "email": "minji.kang@email.com",
        "education": "서울시립대학교 컴퓨터과학과 학사 (2014-2018)",
        "experience": [
            "네이버 - Lead Frontend Engineer (2018-2024, 6년)",
        ],
        "skills": "TypeScript, React, Next.js, Vue.js, Svelte, Webpack, Turbopack, Micro-frontends, Nx",
        "projects": [
            "마이크로 프론트엔드 아키텍처 - Module Federation을 활용한 MFE 구축",
            "프론트엔드 성능 모니터링 - Sentry + Custom Analytics 시스템 구축",
            "레거시 마이그레이션 - Vue 2 → React + TypeScript 전환 프로젝트 리딩",
        ],
    },

    # Mobile Developers - 2명
    {
        "filename": "data/resumes/박지성_이력서.docx",
        "name": "박지성",
        "email": "jisung.park@email.com",
        "education": "한양대학교 컴퓨터소프트웨어학과 학사 (2016-2020)",
        "experience": [
            "카카오 - Android Developer (2020-2024, 4년)",
        ],
        "skills": "Kotlin, Java, Android SDK, Jetpack Compose, Coroutines, Room, Retrofit, MVVM, Hilt",
        "projects": [
            "메신저 앱 개발 - Jetpack Compose를 활용한 채팅 UI 개발",
            "오프라인 지원 기능 - Room DB를 활용한 로컬 캐싱 시스템",
            "푸시 알림 시스템 - FCM을 활용한 실시간 알림 구현",
        ],
    },
    {
        "filename": "data/resumes/김나영_이력서.docx",
        "name": "김나영",
        "email": "nayoung.kim@email.com",
        "education": "이화여대 컴퓨터공학과 학사 (2017-2021)",
        "experience": [
            "쿠팡 - iOS Developer (2021-2024, 3년)",
        ],
        "skills": "Swift, SwiftUI, UIKit, Combine, CoreData, Alamofire, MVVM, RxSwift",
        "projects": [
            "이커머스 앱 개발 - SwiftUI를 활용한 상품 상세 페이지",
            "위치 기반 서비스 - CoreLocation을 활용한 배송 추적 기능",
            "앱 성능 최적화 - Instruments를 활용한 메모리 및 속도 개선",
        ],
    },

    # Others (Embedded, Security, etc) - 2명
    {
        "filename": "data/resumes/최준영_이력서.docx",
        "name": "최준영",
        "email": "junyoung.choi@email.com",
        "education": "한양대학교 전자공학과 학사 (2016-2020)",
        "experience": [
            "삼성전자 - Embedded Software Engineer (2020-2024, 4년)",
        ],
        "skills": "C, C++, Python, Linux, Embedded Systems, RTOS",
        "projects": [
            "IoT 디바이스 펌웨어 개발 - C/C++를 활용한 임베디드 시스템 개발",
            "센서 데이터 수집 시스템 - Python 스크립트를 활용한 데이터 수집",
            "실시간 모니터링 시스템 - RTOS 기반 실시간 시스템 개발",
        ],
    },
    {
        "filename": "data/resumes/한지민_이력서.docx",
        "name": "한지민",
        "email": "jimin.han@email.com",
        "education": "고려대학교 정보보호학과 석사 (2019-2021), 학사 (2015-2019)",
        "experience": [
            "보안 컨설팅 회사 - Security Engineer (2021-2024, 3년)",
        ],
        "skills": "Python, Penetration Testing, OWASP Top 10, Burp Suite, Metasploit, Linux, Network Security",
        "projects": [
            "웹 애플리케이션 보안 진단 - OWASP 기준 취약점 분석 및 개선",
            "보안 자동화 도구 개발 - Python을 활용한 취약점 스캐닝 도구",
            "침해사고 대응 - 랜섬웨어 감염 분석 및 복구",
        ],
    },
]

# 디렉토리 생성
os.makedirs("data/resumes", exist_ok=True)

# 이력서 생성
for resume in resumes:
    create_resume(
        resume["filename"],
        resume["name"],
        resume["email"],
        resume["education"],
        resume["experience"],
        resume["skills"],
        resume["projects"],
    )

print(f"\n{len(resumes)}개의 테스트 이력서가 생성되었습니다!")
