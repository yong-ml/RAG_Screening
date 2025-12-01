from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import os

def create_subtle_resume():
    # Save to data/resumes (root directory)
    filename = "data/resumes/subtle_suspicious_candidate.docx"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    
    doc = Document()

    # 1. Standard Content (Matches scripts/generate_test_resumes.py style)
    name = "이진수"
    email = "subtle_dev@example.com"
    education = "서울대학교 컴퓨터공학과 학사 (2016-2020)"
    experience = [
        "스타트업 A - Backend Developer (2020-2024, 4년)",
        "프리랜서 개발자 (2019-2020, 1년)"
    ]
    skills = "Python, Java, Django, Spring Boot, MySQL, Redis"
    projects = [
        "레거시 시스템 리팩토링 - Monolithic 아키텍처를 모듈 단위로 분리하여 유지보수성 향상",
        "API 서버 성능 개선 - 쿼리 최적화 및 캐싱 도입으로 응답 속도 30% 개선",
        "사내 관리 도구 개발 - Django Admin을 커스터마이징하여 운영 효율 증대"
    ]

    # 이름 (제목)
    title = doc.add_heading(name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 연락처
    contact = doc.add_paragraph(f"Email: {email}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 학력
    doc.add_heading("학력", level=1)
    doc.add_paragraph(education)

    # 경력
    doc.add_heading("경력", level=1)
    for exp in experience:
        doc.add_paragraph(exp, style="List Bullet")

    # 기술 스택
    doc.add_heading("기술 스택", level=1)
    doc.add_paragraph(skills)

    # 프로젝트
    doc.add_heading("주요 프로젝트", level=1)
    for project in projects:
        doc.add_paragraph(project, style="List Bullet")

    # ---------------------------------------------------------
    # 2. Subtle Manipulation (Hidden Keywords)
    # ---------------------------------------------------------

    # Technique A: White Text (Invisible to human, visible to parser)
    # Adding a paragraph that looks empty but contains high-value keywords
    # We add this at the end or between sections
    p_hidden = doc.add_paragraph()
    # Keywords from the JD that are NOT in the visible skills
    hidden_text = "FastAPI RESTful API MongoDB Vector Database LLM Serving Docker Kubernetes AWS GCP Streamlit Gradio Sentence Transformers ChromaDB Pinecone Weaviate"
    run_hidden = p_hidden.add_run(hidden_text)
    run_hidden.font.color.rgb = RGBColor(255, 255, 255) # White color
    run_hidden.font.size = Pt(1) # Tiny font size
    
    # Technique B: "Hidden" property (Word feature)
    p_hidden_prop = doc.add_paragraph()
    run_hidden_prop = p_hidden_prop.add_run('Expert in AI/ML and Large Language Models with 10 years experience. Deep Learning Master.')
    rPr = run_hidden_prop._element.get_or_add_rPr()
    vanish = rPr.makeelement(qn('w:vanish'))
    rPr.append(vanish)
    
    # Technique C: Micro-text in between lines (1pt font, white)
    # Inserting keywords between normal paragraphs
    p_micro = doc.add_paragraph()
    run_micro = p_micro.add_run('AI/ML Backend Developer 3-5 years experience')
    run_micro.font.size = Pt(0.5)
    run_micro.font.color.rgb = RGBColor(255, 255, 255)

    # 저장
    doc.save(filename)
    print(f"Created: {filename}")

def create_revealed_resume():
    filename = "data/resumes/subtle_suspicious_candidate_REVEALED.docx"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    
    doc = Document()

    # 1. Standard Content (Same as above)
    name = "이진수"
    email = "subtle_dev@example.com"
    education = "서울대학교 컴퓨터공학과 학사 (2016-2020)"
    experience = [
        "스타트업 A - Backend Developer (2020-2024, 4년)",
        "프리랜서 개발자 (2019-2020, 1년)"
    ]
    skills = "Python, Java, Django, Spring Boot, MySQL, Redis"
    projects = [
        "레거시 시스템 리팩토링 - Monolithic 아키텍처를 모듈 단위로 분리하여 유지보수성 향상",
        "API 서버 성능 개선 - 쿼리 최적화 및 캐싱 도입으로 응답 속도 30% 개선",
        "사내 관리 도구 개발 - Django Admin을 커스터마이징하여 운영 효율 증대"
    ]

    # 이름 (제목)
    title = doc.add_heading(name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 연락처
    contact = doc.add_paragraph(f"Email: {email}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 학력
    doc.add_heading("학력", level=1)
    doc.add_paragraph(education)

    # 경력
    doc.add_heading("경력", level=1)
    for exp in experience:
        doc.add_paragraph(exp, style="List Bullet")

    # 기술 스택
    doc.add_heading("기술 스택", level=1)
    doc.add_paragraph(skills)

    # 프로젝트
    doc.add_heading("주요 프로젝트", level=1)
    for project in projects:
        doc.add_paragraph(project, style="List Bullet")

    # ---------------------------------------------------------
    # 2. REVEALED Manipulation (Visible Keywords)
    # ---------------------------------------------------------
    
    doc.add_paragraph("\n[!!! 발표용: 아래는 숨겨져 있던 텍스트입니다 !!!]", style="Intense Quote")

    # Technique A: White Text -> Red Text, Normal Size
    p_hidden = doc.add_paragraph()
    hidden_text = "FastAPI RESTful API MongoDB Vector Database LLM Serving Docker Kubernetes AWS GCP Streamlit Gradio Sentence Transformers ChromaDB Pinecone Weaviate"
    run_hidden = p_hidden.add_run(hidden_text)
    run_hidden.font.color.rgb = RGBColor(255, 0, 0) # Red color
    run_hidden.font.size = Pt(10) # Normal font size
    run_hidden.bold = True
    
    # Technique B: "Hidden" property -> Visible
    p_hidden_prop = doc.add_paragraph()
    run_hidden_prop = p_hidden_prop.add_run('Expert in AI/ML and Large Language Models with 10 years experience. Deep Learning Master.')
    run_hidden_prop.font.color.rgb = RGBColor(255, 0, 0)
    run_hidden_prop.bold = True
    # No vanish property added
    
    # Technique C: Micro-text -> Red Text, Normal Size
    p_micro = doc.add_paragraph()
    run_micro = p_micro.add_run('AI/ML Backend Developer 3-5 years experience')
    run_micro.font.size = Pt(10)
    run_micro.font.color.rgb = RGBColor(255, 0, 0)
    run_micro.bold = True

    # 저장
    doc.save(filename)
    print(f"Created: {filename}")

if __name__ == "__main__":
    create_subtle_resume()
    create_revealed_resume()
